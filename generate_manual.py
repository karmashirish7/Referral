"""
Referral Network Portal — Manual Generator (Supabase Edition)
Run:  pip install python-docx
      python3 generate_manual.py
Output: MANUAL.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

NAVY  = RGBColor(0x1F, 0x38, 0x64)
GREEN = RGBColor(0x05, 0x96, 0x69)
RED   = RGBColor(0xDC, 0x26, 0x26)

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ──────────────────────────────────────────────────────

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = NAVY
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    if p.runs: p.runs[0].font.size = Pt(11)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    return p

def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x05, 0x6b, 0x6b)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(2)
    return p

def note(text, colour=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("⚠  " + text)
    r.font.size = Pt(10); r.italic = True
    r.font.color.rgb = colour or RGBColor(0xD9, 0x77, 0x06)
    return p

def tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("✅  " + text)
    r.font.size = Pt(10); r.italic = True
    r.font.color.rgb = GREEN
    return p

def divider():
    doc.add_paragraph("─" * 90)

def table2(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers): hdr[i].text = h
    for row in rows:
        r = t.add_row().cells
        for i, v in enumerate(row): r[i].text = str(v)
    return t

# ════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(48)
r = title.add_run("Referral Network Portal")
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Administrator & Developer Manual  —  Supabase Edition")
r2.font.size = Pt(14); r2.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

doc.add_paragraph()
dp = doc.add_paragraph()
dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
dp.add_run(f"Version 2.0  ·  Generated {datetime.date.today().strftime('%d %B %Y')}").font.size = Pt(10)

doc.add_paragraph()
tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = tag.add_run("Australian Mortgage & Real Estate Sector  ·  Industry Capstone Project")
r3.font.size = Pt(10); r3.italic = True; r3.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════
h1("Table of Contents")
toc = [
    ("1",  "Overview & Key Features"),
    ("2",  "System Requirements"),
    ("3",  "Setting Up Supabase (Database)"),
    ("4",  "Installation on Any Computer"),
    ("5",  "Default Login Accounts"),
    ("6",  "File Structure"),
    ("7",  "User Roles & Permissions"),
    ("8",  "Changing the Primary Colour"),
    ("9",  "Adding / Modifying Referral Statuses"),
    ("10", "Commission Rates & Tiers"),
    ("11", "Adding New Loan Types"),
    ("12", "Changing Database Credentials"),
    ("13", "File Upload Settings"),
    ("14", "Adding New Pages"),
    ("15", "Email Notifications"),
    ("16", "Compliance Notes (Privacy Act 1988 / Spam Act 2003)"),
    ("17", "Troubleshooting"),
]
for num, t in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run(f"  {num}.  {t}").font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ════════════════════════════════════════════════════════════════
h1("1. Overview & Key Features")
body("The Referral Network Portal is a web application that manages partner referrals for Australian mortgage brokerages. It replaces spreadsheets and email tracking with a centralised digital system compliant with the Privacy Act 1988.")
body("The system uses PHP for the backend and Supabase (PostgreSQL) as the cloud database — meaning the data is always online and accessible from any computer without local database setup.")

h2("Key Features")
features = [
    "Partner dashboard — live stats, monthly referral chart, commission summary",
    "6-stage referral tracking: Pending → Qualified → Lodged → Approved → Settled → Declined",
    "Automatic commission calculation: Broker Upfront × Tier Rate (Gold 25%, Silver 20%, Bronze 15%)",
    "Broker Kanban pipeline view (6 columns)",
    "Admin panel — approve users, manage tiers, override commissions",
    "Auditor view — full read-only audit log with CSV export",
    "In-app notifications on every status change",
    "File upload (PDF, DOCX, XLSX) attached to referrals",
    "Role-based access: Partner, Broker, Admin, Auditor",
    "Privacy Act 1988 compliant — client consent checkbox with timestamp, 7-year audit trail",
    "Cloud database via Supabase — works from any computer, no local MySQL needed",
]
for f in features: bullet(f)

# ════════════════════════════════════════════════════════════════
# 2. SYSTEM REQUIREMENTS
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("2. System Requirements")
body("Because the database is hosted on Supabase (cloud), you only need PHP and an internet connection on each computer.")

table2(
    ["Requirement", "Details"],
    [
        ["PHP",            "Version 8.1 or higher (8.2 recommended). Must include pdo_pgsql extension."],
        ["Web server",     "PHP built-in server (php -S) for development, or Apache/Nginx for production."],
        ["Supabase account","Free account at supabase.com — one project shared by all computers."],
        ["Internet",       "Required to reach the Supabase cloud database."],
        ["Browser",        "Chrome, Firefox, Safari, or Edge (latest version)."],
        ["No local MySQL", "MySQL is NOT needed. Supabase replaces it entirely."],
    ]
)

h2("Checking pdo_pgsql is installed")
body("Run this in a terminal to confirm the PostgreSQL driver is available:")
code("php -m | grep pgsql")
body("You should see pgsql and pdo_pgsql in the output. If not:")
code("# Mac (Homebrew):    brew install php")
code("# Ubuntu/Debian:     sudo apt install php-pgsql")
code("# Windows (XAMPP):   enable extension=pdo_pgsql in php.ini")

# ════════════════════════════════════════════════════════════════
# 3. SETTING UP SUPABASE
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("3. Setting Up Supabase (Database)  —  Do This Once")
body("Supabase is a free cloud PostgreSQL database. You only set this up once; all computers then connect to the same database.")

h2("Step 1 — Create a Supabase account and project")
for s in [
    "Go to https://supabase.com and sign up for a free account.",
    'Click "New Project". Choose a name (e.g. "ReferralPortal") and set a strong database password. Save this password — you will need it.',
    "Wait ~2 minutes for the project to be created.",
]:
    bullet(s)

h2("Step 2 — Import the database schema")
for s in [
    'In your Supabase project, click "SQL Editor" in the left sidebar.',
    'Click "New Query".',
    "Open the file database.sql from this project folder. Select all the text (Ctrl+A / Cmd+A) and copy it.",
    'Paste it into the SQL Editor and click "Run".',
    'You should see "Success. No rows returned." — this means all tables and sample data were created.',
]:
    bullet(s)
note("Only run database.sql once. Running it again is safe (it uses ON CONFLICT DO NOTHING) but will not re-insert data that already exists.")

h2("Step 3 — Find your connection details")
body("Go to: Supabase Dashboard → Project Settings (gear icon) → Database → Connection parameters")
table2(
    ["Field", "Where to find it", "Example value"],
    [
        ["Host",     "Connection parameters → Host",     "db.abcdefghijkl.supabase.co"],
        ["Port",     "Always 5432",                       "5432"],
        ["Database", "Always 'postgres'",                 "postgres"],
        ["User",     "Always 'postgres'",                 "postgres"],
        ["Password", "The password you set in Step 1",    "your-secret-password"],
    ]
)
tip("Copy the Host value carefully — it is a long string like db.abcdefghijkl.supabase.co")

# ════════════════════════════════════════════════════════════════
# 4. INSTALLATION ON ANY COMPUTER
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("4. Installation on Any Computer")
body("Every computer that runs the portal follows the same 4 steps. The database is already set up in Supabase — you just need to connect to it.")

h2("Step 1 — Get the project files")
body("Option A — Clone from GitHub (recommended):")
code("git clone https://github.com/karmashirish7/Referral")
code("cd Referral")
body("Option B — Download the ZIP from GitHub and extract it.")

h2("Step 2 — Start the PHP web server")
code("php -S localhost:8080")
body("Keep this terminal window open while using the portal. To stop the server, press Ctrl+C.")
note("If port 8080 is in use, change it: php -S localhost:9000")

h2("Step 3 — Run the installer")
for s in [
    "Open a browser and go to:  http://localhost:8080/install.php",
    "The Setup Wizard will appear. Enter your Supabase connection details from Section 3.",
    'Click "Connect & Save".',
    "If successful, you will see a green confirmation screen.",
]:
    bullet(s)

h2("Step 4 — Log in")
body("Go to http://localhost:8080 and log in with any of the default accounts (see Section 5).")
tip("The installer creates a config.php file locally with your credentials. This file is in .gitignore so it is never uploaded to GitHub — each computer keeps its own copy.")

# ════════════════════════════════════════════════════════════════
# 5. DEFAULT LOGIN ACCOUNTS
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("5. Default Login Accounts")
note("All default passwords are:  password   Change them immediately before real use!")

table2(
    ["Name", "Email", "Role", "Access Level"],
    [
        ["Admin User",    "admin@networkportal.com",  "Admin",   "Full system access — users, commissions, audit log"],
        ["James Broker",  "broker@networkportal.com", "Broker",  "Pipeline, status updates, broker notes"],
        ["Alex Thompson", "alex@partner.com",         "Partner", "Submit referrals, view own commissions"],
        ["Sarah Chen",    "sarah@partner.com",        "Partner", "Submit referrals, view own commissions"],
        ["Lisa Audit",    "auditor@networkportal.com","Auditor", "Read-only audit log and referral history"],
    ]
)

h2("How to change a password")
body("Log in → click Settings in the sidebar → Change Password section. Enter your current password and a new one (minimum 8 characters).")

h2("How to reset a password via Supabase SQL Editor")
body("If you are locked out, go to Supabase → SQL Editor and run:")
code("UPDATE users")
code("SET password = '$2y$10$92IXUNpkjO0rOQ5byMi.Ye4oKoEa3Ro9llC/.og/at2.uheWG/igi'")
code("WHERE email = 'your@email.com';")
body("This resets the password back to:  password")
body("Then log in and change it immediately via Settings.")

# ════════════════════════════════════════════════════════════════
# 6. FILE STRUCTURE
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("6. File Structure")
table2(
    ["File / Folder", "Purpose"],
    [
        ("index.php",               "Login page — role selector, email, password"),
        ("register.php",            "New partner/broker registration form"),
        ("logout.php",              "Destroys session and redirects to login"),
        ("dashboard.php",           "Partner dashboard — stats, chart, pipeline, payouts"),
        ("referrals.php",           "Referral list for all roles — filters, CSV export"),
        ("submit-referral.php",     "New referral form — client details, consent, file upload"),
        ("pipeline.php",            "Broker Kanban board — 6 stage columns"),
        ("admin-users.php",         "Admin: approve/suspend users, add user, set tier"),
        ("admin-commissions.php",   "Admin: mark commissions paid, override amounts"),
        ("audit.php",               "Admin/Auditor: full audit trail with filters and CSV export"),
        ("documents.php",           "View and upload files attached to referrals"),
        ("settings.php",            "Update profile, change password, email preferences"),
        ("notifications.php",       "View all in-app notifications"),
        ("install.php",             "Supabase setup wizard — run once per computer"),
        ("database.sql",            "PostgreSQL schema — paste into Supabase SQL Editor once"),
        ("config.php",              "Auto-created by installer. Contains DB credentials. NEVER commit to git."),
        ("includes/db.php",         "Database connection using pdo_pgsql driver"),
        ("includes/auth.php",       "Session helpers: isLoggedIn(), requireRole(), etc."),
        ("includes/functions.php",  "Helpers: commission calc, audit log, notifications, badges"),
        ("includes/sidebar.php",    "Role-aware left sidebar component"),
        ("assets/css/style.css",    "All custom styles — colours, layout, cards, badges"),
        ("assets/js/main.js",       "Notifications dropdown, file upload drag-drop, confirmations"),
        ("uploads/",                "Uploaded referral documents stored here (local only)"),
        ("generate_manual.py",      "Run this to regenerate MANUAL.docx after making changes"),
    ]
)

note("config.php is in .gitignore and must never be committed. Each computer creates its own copy via install.php.")

# ════════════════════════════════════════════════════════════════
# 7. USER ROLES & PERMISSIONS
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("7. User Roles & Permissions")

table2(
    ["Role", "Can Do", "Cannot Do"],
    [
        ["Partner",
         "Submit referrals · View own referrals · View own commissions & payouts · Upload documents · Update profile",
         "See other partners' referrals · Change referral status · Access admin or audit sections"],
        ["Broker",
         "View all referrals (assigned to them) · Update referral status · Enter broker upfront commission · Add broker notes · Kanban pipeline",
         "Approve users · Override commissions · View audit log"],
        ["Admin",
         "Everything above · Approve/suspend user accounts · Add new users · Change partner tier · Override commissions · Mark commissions paid · View full audit log",
         "Nothing — full access"],
        ["Auditor",
         "Read-only view of audit log · Read-only view of all referrals",
         "Cannot change any data — read-only access only"],
    ]
)

h2("Referral Status Flow")
body("The six statuses and who can change them:")
table2(
    ["Status", "Meaning", "Changed By"],
    [
        ["Pending",   "Newly submitted, awaiting broker review",         "Broker / Admin"],
        ["Qualified", "Broker has spoken to client, lead is valid",       "Broker / Admin"],
        ["Lodged",    "Finance application submitted to lender",          "Broker / Admin"],
        ["Approved",  "Lender has approved the application",              "Broker / Admin"],
        ["Settled",   "Loan settled — commission auto-calculated",        "Broker / Admin"],
        ["Declined",  "Lead did not proceed",                             "Broker / Admin"],
    ]
)
tip("When status is set to Settled, the broker must enter the Broker Upfront Commission amount. The system automatically calculates the partner's commission and creates a commission record.")

# ════════════════════════════════════════════════════════════════
# 8. CHANGING THE PRIMARY COLOUR
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("8. Changing the Primary Colour")
body("The portal uses dark navy #1F3864 as defined in the Industry Project spec. To change it, open assets/css/style.css and edit the three CSS variables at the very top of the file:")
code(":root {")
code("    --navy:       #1F3864;   /* sidebar, buttons, headings */")
code("    --navy-dark:  #172d55;   /* button hover state        */")
code("    --navy-light: #2a4a7f;   /* active nav item           */")
code("}")
body("Change all three values consistently to your desired colour, save the file, and refresh the browser.")
note("The monthly bar chart in dashboard.php also uses the colour inline. Search for '#1F3864' in dashboard.php and update those values too.")

# ════════════════════════════════════════════════════════════════
# 9. ADDING / MODIFYING REFERRAL STATUSES
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("9. Adding / Modifying Referral Statuses")
body("Statuses must be updated in three places consistently.")

h2("9a — Update the database (Supabase SQL Editor)")
body("Run this in Supabase → SQL Editor to add a new status (e.g. 'on_hold'):")
code("ALTER TABLE referrals")
code("  DROP CONSTRAINT IF EXISTS referrals_status_check;")
code("")
code("ALTER TABLE referrals")
code("  ADD CONSTRAINT referrals_status_check")
code("  CHECK (status IN ('pending','qualified','lodged','approved','settled','declined','on_hold'));")

h2("9b — Update includes/functions.php — statusBadge()")
body("Find the $map array inside statusBadge() and add the new status with a CSS class:")
code("$map = [")
code("    ...existing entries...,")
code("    'on_hold' => 'badge-lodged',  // reuse an existing colour")
code("];")

h2("9c — Add a CSS class (optional, assets/css/style.css)")
body("If you want a unique colour for the new status, add it in the STATUS BADGES section:")
code(".badge-on_hold { background: #e0f2fe; color: #0369a1; }")

h2("9d — Update status dropdowns")
body("Search all PHP files for the hardcoded status option lists and add the new status. Files affected: referrals.php, pipeline.php.")

# ════════════════════════════════════════════════════════════════
# 10. COMMISSION RATES & TIERS
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("10. Commission Rates & Tiers")

h2("Commission formula")
body("When a referral is settled, the partner commission is calculated as:")
code("Partner Commission  =  Broker Upfront Commission  ×  (Tier Rate  ÷  100)")
body("Example: Broker upfront = $8,500 · Partner tier = Gold (25%) → Commission = $2,125")

h2("Default tier rates")
table2(
    ["Tier", "Default Rate", "Typical Partner Type"],
    [
        ["Gold",   "25%", "High-volume or long-standing partners"],
        ["Silver", "20%", "Standard partners (default on registration)"],
        ["Bronze", "15%", "New or lower-volume partners"],
    ]
)

h2("Where to change default rates in code")
body("The tier-to-rate mapping appears in three files. Update all three with the same values:")
code("admin-users.php  (PHP):  $rate = ['Gold'=>25,'Silver'=>20,'Bronze'=>15][$tier] ?? 20;")
code("admin-users.php  (JS):   const tierRates = {Gold:25, Silver:20, Bronze:15};")
code("register.php     (PHP):  $rate = ['Gold'=>25,'Silver'=>20,'Bronze'=>15][$tier] ?? 20;")

h2("Changing a single user's rate")
body("Log in as Admin → Users → click Edit Tier next to a partner → enter a custom rate in the Custom Commission Rate % field. This overrides the tier default for that user only.")

h2("Commission calculation function")
body("The formula lives in includes/functions.php — function calculateCommission():")
code("function calculateCommission($brokerUpfront, $tierRate) {")
code("    return round($brokerUpfront * ($tierRate / 100), 2);")
code("}")
body("Modify this if the formula changes (e.g. add a cap, deduct a flat fee).")

# ════════════════════════════════════════════════════════════════
# 11. ADDING NEW LOAN TYPES
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("11. Adding New Loan Types")
body("Loan types are stored as a CHECK constraint in PostgreSQL and as a PHP array in the form.")

h2("Step 1 — Update the database constraint (Supabase SQL Editor)")
code("ALTER TABLE referrals")
code("  DROP CONSTRAINT IF EXISTS referrals_loan_type_check;")
code("")
code("-- Add your new type to this list:")
code("ALTER TABLE referrals")
code("  ADD CONSTRAINT referrals_loan_type_check")
code("  CHECK (loan_type IN ('Owner-Occupied','Investment','Refinance','Commercial','Construction','SMSF'));")

h2("Step 2 — Update submit-referral.php")
body("Find the $loanTypes array near the top of the file and add your new type:")
code("$loanTypes = ['Owner-Occupied','Investment','Refinance','Commercial','Construction','SMSF'];")

# ════════════════════════════════════════════════════════════════
# 12. CHANGING DATABASE CREDENTIALS
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("12. Changing Database Credentials")
body("Credentials are stored in config.php (auto-created by install.php). To change them, either re-run the installer or edit config.php directly:")
code("<?php")
code("define('DB_HOST', 'db.xxxxxxxxxxxx.supabase.co');")
code("define('DB_PORT', '5432');")
code("define('DB_USER', 'postgres');")
code("define('DB_PASS', 'your-supabase-password');")
code("define('DB_NAME', 'postgres');")

note("config.php is in .gitignore and must NEVER be committed to GitHub. It contains your database password.")

h2("Re-running the installer")
body("To reconfigure, visit:  http://localhost:8080/install.php?reinstall=1")
body("This overwrites config.php with the new credentials.")

h2("Rotating your Supabase password")
for s in [
    "Go to Supabase Dashboard → Project Settings → Database → Reset database password.",
    "Copy the new password.",
    "On each computer, run install.php?reinstall=1 or edit config.php with the new password.",
]:
    bullet(s)

# ════════════════════════════════════════════════════════════════
# 13. FILE UPLOAD SETTINGS
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("13. File Upload Settings")

h2("Allowed file types")
body("Currently: PDF, DOCX, DOC, XLSX, XLS. To change, edit the $allowed array in both submit-referral.php and documents.php:")
code("$allowed = ['pdf','docx','doc','xlsx','xls'];")

h2("Maximum file size")
body("The app enforces 10MB. Change $maxSize in submit-referral.php:")
code("$maxSize = 10 * 1024 * 1024; // change 10 to desired MB")
note("PHP also has its own upload limit in php.ini. You may need to increase upload_max_filesize and post_max_size if you raise the limit above 8MB.")

h2("Upload folder")
body("Files are saved to the uploads/ folder in the project directory. This folder is local — files are NOT stored in Supabase. Every computer has its own local uploads/ folder.")
note("If you need files to be shared across computers, consider using Supabase Storage (supabase.com/docs/guides/storage) and updating the upload code in submit-referral.php and documents.php.")

# ════════════════════════════════════════════════════════════════
# 14. ADDING NEW PAGES
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("14. Adding New Pages")
body("Copy this skeleton into a new .php file in the project root:")
code("<?php")
code("require_once 'includes/db.php';")
code("require_once 'includes/auth.php';")
code("require_once 'includes/functions.php';")
code("requireRole('partner'); // or requireLogin() for any role")
code("")
code("$user   = currentUser();")
code("$uid    = currentUserId();")
code("$unread = unreadNotificationCount($pdo, $uid);")
code("?>")
code("<!DOCTYPE html><html lang=\"en\"><head>")
code("    <link rel=\"stylesheet\" href=\"https://cdn.jsdelivr.net/npm/bootstrap-icons@1.11.0/font/bootstrap-icons.css\">")
code("    <link rel=\"stylesheet\" href=\"assets/css/style.css\">")
code("</head><body>")
code("<div class=\"app-wrapper\">")
code("    <?php include 'includes/sidebar.php'; ?>")
code("    <div class=\"main-content\">")
code("        <!-- copy any top-header block from another page -->")
code("        <div class=\"page-body\">")
code("            <div class=\"page-heading\">")
code("                <div><h1>Your Page Title</h1></div>")
code("            </div>")
code("            <!-- your content here -->")
code("        </div>")
code("    </div>")
code("</div>")
code("<script src=\"assets/js/main.js\"></script>")
code("</body></html>")

h2("Adding a sidebar link")
body("Open includes/sidebar.php, find the role section you want, and add:")
code("<?= navItem('your-page.php', 'icon-name', 'Link Label', 'your-page.php', $page) ?>")
body("Browse Bootstrap icon names at:  https://icons.getbootstrap.com")

# ════════════════════════════════════════════════════════════════
# 15. EMAIL NOTIFICATIONS
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("15. Email Notifications")
body("The portal currently sends in-app notifications only. Real email delivery requires an SMTP service or transactional email provider.")

h2("Recommended providers (free tiers available)")
table2(
    ["Provider", "Free Tier", "How to integrate"],
    [
        ["Resend",     "3,000 emails/month", "API key + curl or SDK in PHP"],
        ["Mailgun",    "100 emails/day",     "SMTP or API"],
        ["SendGrid",   "100 emails/day",     "SMTP or API"],
        ["PHPMailer",  "SMTP (your account)","Composer: composer require phpmailer/phpmailer"],
    ]
)

h2("Where to add email sending")
body("Open includes/functions.php. The addNotification() function is called every time a notification is created. Add your email code inside it:")
code("function addNotification($pdo, $userId, $title, $message) {")
code("    // existing DB insert (keep this) ...")
code("")
code("    // Add email sending below:")
code("    $stmt = $pdo->prepare('SELECT email, consent_email FROM users WHERE id=?');")
code("    $stmt->execute([$userId]);")
code("    $u = $stmt->fetch();")
code("    if ($u && $u['consent_email']) {")
code("        // mail($u['email'], $title, $message); // basic PHP mail()")
code("        // or use PHPMailer / Resend API here")
code("    }")
code("}")
note("Only send emails when consent_email = 1 to comply with the Spam Act 2003.")

# ════════════════════════════════════════════════════════════════
# 16. COMPLIANCE NOTES
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("16. Compliance Notes  —  Privacy Act 1988 & Spam Act 2003")

h2("Client Consent (APP 3)")
body("Every referral submission includes a mandatory consent checkbox. The timestamp of consent is stored in the consent_timestamp column of the referrals table. This field must not be removed.")

h2("Audit Trail (APP 11)")
body("Every significant action — login, logout, referral submission, status change, commission payment, user approval — is written to the audit_log table via the logAction() function in includes/functions.php. The audit log is accessible to Admin and Auditor roles at audit.php and can be exported as CSV.")

h2("Data Retention (APP 11.2)")
body("The Privacy Act 1988 requires personal data to be kept for 7 years then destroyed or de-identified. The system does not auto-delete records. Implement a manual or scheduled process to anonymise records older than 7 years.")
body("Anonymisation query example (run in Supabase SQL Editor):")
code("UPDATE referrals")
code("SET client_name='[Anonymised]', client_email=NULL, client_phone=NULL")
code("WHERE date_submitted < NOW() - INTERVAL '7 years';")

h2("Email Opt-Out (Spam Act 2003)")
body("Users can unsubscribe from email notifications in Settings → uncheck 'Receive email notifications'. This sets consent_email = 0. Email code must check this field before sending (already built into addNotification()).")

h2("Password Security")
body("All passwords are stored as bcrypt hashes using PHP's password_hash(PASSWORD_DEFAULT). Plain-text passwords are never stored or logged.")

h2("Data in Transit")
body("The Supabase connection uses sslmode=require, encrypting all data between the PHP server and the database. Use HTTPS in production (certificate from Let's Encrypt or your hosting provider).")

# ════════════════════════════════════════════════════════════════
# 17. TROUBLESHOOTING
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("17. Troubleshooting")

issues = [
    (
        "Browser shows 'ERR_CONNECTION_REFUSED'",
        "The PHP server is not running. Open a terminal in the project folder and run:\n  php -S localhost:8080\nKeep the terminal open while using the portal."
    ),
    (
        "Installer shows 'pdo_pgsql driver not found'",
        "The PHP PostgreSQL extension is missing.\n"
        "  Mac (Homebrew):  brew install php\n"
        "  Ubuntu/Debian:   sudo apt install php-pgsql\n"
        "  Windows (XAMPP): open php.ini, uncomment extension=pdo_pgsql, restart Apache."
    ),
    (
        "Installer shows 'could not connect to server'",
        "Check the Host field — it must be the full Supabase host like db.abcdefghijkl.supabase.co\n"
        "Also check your internet connection, and that the Supabase project is not paused (free projects pause after 1 week of inactivity — unpause in the Supabase dashboard)."
    ),
    (
        "Installer shows 'wrong password'",
        "Re-check the database password from Supabase Dashboard → Project Settings → Database.\n"
        "Note: this is the DATABASE password, not your Supabase login password."
    ),
    (
        "'Pending admin approval' message on login",
        "The account has status = 'pending'. Log in as Admin, go to admin-users.php, and click Approve next to the user."
    ),
    (
        "Supabase project is paused",
        "Free Supabase projects pause after approximately 1 week of inactivity. Go to supabase.com, open your project, and click 'Restore project'. It takes about 1 minute to resume."
    ),
    (
        "Commission not calculated when settling",
        "The broker must enter the Broker Upfront Commission amount in the status update modal when changing status to 'Settled'. Without that value, no commission record is created."
    ),
    (
        "File upload fails",
        "Check the uploads/ folder exists and is writable (chmod 775 uploads/ on Mac/Linux).\n"
        "Also check PHP's upload_max_filesize in php.ini is at least 10M."
    ),
    (
        "Dashboard chart not showing",
        "Chart.js is loaded from CDN — check your internet connection. Alternatively, download chart.umd.min.js and reference it locally in dashboard.php."
    ),
    (
        "Icons showing as empty boxes",
        "Bootstrap Icons is loaded from CDN. Check internet connection or download the icon font and serve it locally from assets/css/."
    ),
    (
        "White screen / PHP error",
        "Add this to the top of the failing file temporarily:\n"
        "  ini_set('display_errors', 1); error_reporting(E_ALL);\n"
        "This will show the error message. Remove it before using the portal."
    ),
]

for prob, sol in issues:
    h2("Problem: " + prob)
    body("Solution: " + sol)

# ════════════════════════════════════════════════════════════════
# FOOTER
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("Referral Network Portal  —  Confidential Developer Manual")
r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

doc.add_paragraph()
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.add_run(f"© {datetime.date.today().year} Network Portal · Supabase Edition · Professional Use Only").font.size = Pt(9)

doc.save("MANUAL.docx")
print("✅  MANUAL.docx saved successfully.")
